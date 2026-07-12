from __future__ import annotations

import csv
import hashlib
import json
import re
import uuid
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader


def file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def _chunk(text: str, prefix: str, size: int = 1400) -> list[dict]:
    paragraphs = [re.sub(r"\s+", " ", p).strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    result, current = [], ""
    n = 1
    for p in paragraphs:
        if current and len(current) + len(p) > size:
            result.append({"location": f"{prefix}, chunk {n}", "text": current})
            current, n = p, n + 1
        else:
            current = f"{current}\n{p}".strip()
    if current:
        result.append({"location": f"{prefix}, chunk {n}", "text": current})
    return result


def parse_file(path: Path) -> tuple[list[dict], dict]:
    suffix = path.suffix.lower()
    chunks: list[dict] = []
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages, 1):
            chunks.extend(_chunk(page.extract_text() or "", f"page {i}"))
        details = {"pages": len(reader.pages)}
    elif suffix == ".docx":
        doc = Document(str(path))
        heading = "Document start"
        rows = []
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith("Heading"):
                heading = p.text.strip() or heading
            elif p.text.strip():
                rows.append((heading, p.text.strip()))
        for table_no, table in enumerate(doc.tables, 1):
            table_text = "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            rows.append((f"table {table_no}", table_text))
        for n, (loc, text) in enumerate(rows, 1):
            chunks.append({"location": f"{loc}, item {n}", "text": text})
        details = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
    elif suffix in {".xlsx", ".xlsm"}:
        book = pd.ExcelFile(path)
        for sheet in book.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet, dtype=str).fillna("")
            header = " | ".join(map(str, df.columns))
            for idx, row in df.iterrows():
                text = header + "\n" + " | ".join(map(str, row.tolist()))
                chunks.append({"location": f"sheet {sheet}, row {idx + 2}", "text": text})
        details = {"sheets": book.sheet_names}
    elif suffix == ".csv":
        df = pd.read_csv(path, dtype=str).fillna("")
        for idx, row in df.iterrows():
            chunks.append({"location": f"row {idx + 2}", "text": " | ".join(f"{c}={row[c]}" for c in df.columns)})
        details = {"rows": len(df), "columns": list(df.columns)}
    elif suffix in {".json"}:
        obj = json.loads(path.read_text(encoding="utf-8"))
        chunks = _chunk(json.dumps(obj, ensure_ascii=False, indent=2), "JSON")
        details = {"type": type(obj).__name__}
    else:
        chunks = _chunk(path.read_text(encoding="utf-8", errors="replace"), "text")
        details = {"characters": sum(len(c["text"]) for c in chunks)}
    return chunks, details


def ingest(con, study_id: str, path: Path, kind: str, source_class: str = "project") -> dict:
    digest = file_hash(path)
    doc_id = digest[:16]
    chunks, details = parse_file(path)
    con.execute(
        "INSERT OR REPLACE INTO documents(id,study_id,name,kind,sha256,status,version,details) VALUES(?,?,?,?,?,?,?,?)",
        (doc_id, study_id, path.name, kind, digest, "Parsed", "demo-1.0", json.dumps(details)),
    )
    con.execute("DELETE FROM chunks WHERE document_id=?", (doc_id,))
    for item in chunks:
        cid = uuid.uuid4().hex
        con.execute(
            "INSERT INTO chunks(id,study_id,document_id,document,location,source_class,text) VALUES(?,?,?,?,?,?,?)",
            (cid, study_id, doc_id, path.name, item["location"], source_class, item["text"]),
        )
        con.execute("INSERT INTO chunks_fts(id,text) VALUES(?,?)", (cid, item["text"]))
    con.commit()
    return {"id": doc_id, "name": path.name, "chunks": len(chunks), "details": details}
